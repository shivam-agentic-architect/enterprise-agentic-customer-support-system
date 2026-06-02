import logging
import io
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from uuid import UUID

from pypdf import PdfReader
from docx import Document

from ...database.session import get_db
from ...schemas.index import KBArticleResponse
from ...models.index import KnowledgeDocument
from ...services.bedrock_service import bedrock_service
from ...services.opensearch_service import opensearch_service

router = APIRouter(prefix="/knowledge", tags=["Knowledge Base"])

# Custom recursive-style character splitter helper to maintain cohesive word boundaries
def split_text_into_chunks(text: str, chunk_size: int = 1000, chunk_overlap: int = 150) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            # Attempt to find the nearest paragraph break or space to avoid middle-word cuts
            last_newline = text.rfind('\n', start, end)
            if last_newline > start + chunk_size - chunk_overlap:
                end = last_newline + 1
            else:
                last_space = text.rfind(' ', start, end)
                if last_space > start + chunk_size - chunk_overlap:
                    end = last_space + 1
        chunk_content = text[start:end].strip()
        if chunk_content:
            chunks.append(chunk_content)
        start += chunk_size - chunk_overlap
    return chunks

@router.get("/", response_model=List[KBArticleResponse])
async def list_documents(db: AsyncSession = Depends(get_db)):
    query = select(KnowledgeDocument).filter(KnowledgeDocument.is_deleted == False).order_by(KnowledgeDocument.last_sync.desc())
    result = await db.execute(query)
    return result.scalars().all()

@router.get("/search")
async def search_knowledge(query: str, limit: int = 3, db: AsyncSession = Depends(get_db)):
    if not query:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Query parameter is required."
        )
    try:
        # Generate Titan Embeddings dynamically
        embeddings = await bedrock_service.generate_embeddings(query)
        # Search OpenSearch Vector DB
        hits = await opensearch_service.search_vector_database(query, embeddings, limit=limit)
        return hits
    except Exception as e:
        logging.error(f"RAG Vector Search failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"RAG search execution failed: {e}"
        )

@router.post("/upload", response_model=KBArticleResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...), 
    db: AsyncSession = Depends(get_db)
):
    name = file.filename
    # Simple check format
    ext = name.split(".")[-1].lower() if "." in name else "txt"
    if ext not in ["pdf", "docx", "txt"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, 
            detail="Unsupported format type. Only PDF, DOCX or TXT allowed."
        )

    # Read bytes and extract raw text
    try:
        content_bytes = await file.read()
        size_str = f"{(len(content_bytes) / 1024):.1f} KB" if len(content_bytes) < 1024*1024 else f"{(len(content_bytes) / (1024*1024)):.1f} MB"
        
        raw_text = ""
        if ext == "txt":
            raw_text = content_bytes.decode("utf-8", errors="ignore")
        elif ext == "pdf":
            # Real PDF text parsing using PyPDF
            pdf_file = io.BytesIO(content_bytes)
            reader = PdfReader(pdf_file)
            pdf_pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pdf_pages_text.append(text)
            raw_text = "\n".join(pdf_pages_text)
        elif ext == "docx":
            # Real DOCX text parsing using python-docx
            docx_file = io.BytesIO(content_bytes)
            doc = Document(docx_file)
            raw_text = "\n".join([p.text for p in doc.paragraphs])

        # If document contains zero text, fallback or raise warning
        if not raw_text.strip():
            raw_text = f"Empty or un-extractable structural document parsed: {name}."

        # Perform high-fidelity chunking
        chunks = split_text_into_chunks(raw_text, chunk_size=800, chunk_overlap=120)
        if not chunks:
            chunks = [raw_text]

        # Index document summary record in DB first
        snippet = chunks[0][:300] + "..." if len(chunks[0]) > 300 else chunks[0]
        db_doc = KnowledgeDocument(
            title=name,
            category="policy" if "policy" in name.lower() else "sop",
            content=snippet, 
            format=ext,
            size=size_str,
            status="indexed"
        )
        db.add(db_doc)
        await db.flush()

        # Generate Titan Embeddings and index each chunk dynamically in OpenSearch
        for idx, chunk_text in enumerate(chunks):
            # Compute embeddings for each segment
            chunk_embeddings = await bedrock_service.generate_embeddings(chunk_text)
            
            # Index chunk with unique compound key pointing to the database parent doc
            await opensearch_service.index_document(
                document_id=f"{db_doc.id}_chunk_{idx}",
                title=name,
                content=chunk_text,
                embeddings=chunk_embeddings,
                metadata={
                    "parent_id": str(db_doc.id),
                    "url": f"s3://kb-policies/{name}",
                    "chunk_index": idx
                }
            )

        # Update citations count based on chunks size
        db_doc.citations_count = len(chunks)
        db.add(db_doc)
        await db.commit()
        await db.refresh(db_doc)
        return db_doc
    except Exception as e:
        await db.rollback()
        logging.error(f"Failed parsing uploader: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Upload parsing failed: {e}"
        )
